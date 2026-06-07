"""
WriterFlow — Core Layer v3.1
Database · Repositories · Services · Text Utilities · Export
"""
import sqlite3, os, io, re, base64, logging
from contextlib import contextmanager
from typing import Dict, List, Optional, Tuple
from PIL import Image

log = logging.getLogger(__name__)

# ══════════════════════════════════════════════════════
# DATABASE
# ══════════════════════════════════════════════════════

_env = os.environ.get("WRITERFLOW_DB", "").strip()
DB_PATH = _env if _env else ("/tmp/writerflow.db" if os.path.isdir("/tmp") else "writerflow.db")

@contextmanager
def _conn():
    c = sqlite3.connect(DB_PATH)
    c.row_factory = sqlite3.Row
    c.execute("PRAGMA foreign_keys=ON")
    c.execute("PRAGMA journal_mode=WAL")
    try:
        yield c; c.commit()
    except Exception:
        c.rollback(); raise
    finally:
        c.close()

def init_db():
    with _conn() as c:
        c.executescript("""
        CREATE TABLE IF NOT EXISTS books(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL, synopsis TEXT DEFAULT '',
            genre TEXT DEFAULT '', status TEXT DEFAULT 'Planejamento',
            cover_image BLOB, cover_mime TEXT DEFAULT 'image/jpeg',
            word_count INTEGER DEFAULT 0, deleted_at TIMESTAMP,
            last_opened TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP);
        CREATE TABLE IF NOT EXISTS chapters(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            book_id INTEGER NOT NULL, title TEXT NOT NULL,
            content TEXT DEFAULT '', position INTEGER DEFAULT 0,
            word_count INTEGER DEFAULT 0,
            word_goal INTEGER DEFAULT 0,
            scene_marker TEXT DEFAULT '',
            deleted_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(book_id) REFERENCES books(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS characters(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            book_id INTEGER NOT NULL, name TEXT NOT NULL,
            role TEXT DEFAULT '', description TEXT DEFAULT '',
            photo BLOB, photo_mime TEXT DEFAULT 'image/jpeg',
            relationships TEXT DEFAULT '', notes TEXT DEFAULT '',
            age TEXT DEFAULT '', deleted_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(book_id) REFERENCES books(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS locations(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            book_id INTEGER NOT NULL, name TEXT NOT NULL,
            description TEXT DEFAULT '', notes TEXT DEFAULT '',
            deleted_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(book_id) REFERENCES books(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS factions(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            book_id INTEGER NOT NULL, name TEXT NOT NULL,
            description TEXT DEFAULT '', notes TEXT DEFAULT '',
            deleted_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(book_id) REFERENCES books(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS timeline_events(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            book_id INTEGER NOT NULL, title TEXT NOT NULL,
            description TEXT DEFAULT '', date_label TEXT DEFAULT '',
            position INTEGER DEFAULT 0, deleted_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(book_id) REFERENCES books(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS brain_dumps(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            book_id INTEGER, content TEXT NOT NULL,
            tags TEXT DEFAULT '', pinned INTEGER DEFAULT 0,
            deleted_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(book_id) REFERENCES books(id) ON DELETE SET NULL);
        CREATE TABLE IF NOT EXISTS goals(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            goal_type TEXT NOT NULL, target_words INTEGER DEFAULT 0,
            period TEXT DEFAULT 'daily', active INTEGER DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP);
        CREATE TABLE IF NOT EXISTS writing_sessions(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            book_id INTEGER, chapter_id INTEGER,
            words_written INTEGER DEFAULT 0,
            minutes_written INTEGER DEFAULT 0,
            session_date DATE DEFAULT (DATE('now')),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(book_id) REFERENCES books(id) ON DELETE SET NULL);
        CREATE TABLE IF NOT EXISTS kindle_bookmarks(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            book_id INTEGER NOT NULL,
            chapter_id INTEGER NOT NULL,
            note TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP);
        CREATE TABLE IF NOT EXISTS app_settings(key TEXT PRIMARY KEY, value TEXT);
        CREATE INDEX IF NOT EXISTS idx_ch_book ON chapters(book_id);
        CREATE INDEX IF NOT EXISTS idx_char_book ON characters(book_id);
        CREATE INDEX IF NOT EXISTS idx_bd_book ON brain_dumps(book_id);
        CREATE INDEX IF NOT EXISTS idx_sess_date ON writing_sessions(session_date);
        """)
    # migrations — add new columns to existing databases
    _add_col("books","last_opened","TIMESTAMP")
    _add_col("chapters","word_goal","INTEGER DEFAULT 0")
    _add_col("chapters","scene_marker","TEXT DEFAULT ''")
    _add_col("characters","age","TEXT DEFAULT ''")
    _add_col("brain_dumps","pinned","INTEGER DEFAULT 0")
    _add_col("writing_sessions","minutes_written","INTEGER DEFAULT 0")
    for tbl in ("books","chapters","characters","locations","factions","timeline_events","brain_dumps"):
        _add_col(tbl,"deleted_at","TIMESTAMP")
    for tbl in ("locations","factions","timeline_events"):
        _add_col(tbl,"updated_at","TIMESTAMP",backfill="created_at")
    with _conn() as c:
        c.executescript("""
        CREATE INDEX IF NOT EXISTS idx_books_active ON books(status,genre) WHERE deleted_at IS NULL;
        CREATE INDEX IF NOT EXISTS idx_ch_active ON chapters(book_id,position) WHERE deleted_at IS NULL;
        """)

def _add_col(table, col, typ, backfill=None):
    try:
        with _conn() as c:
            existing = {r[1] for r in c.execute(f"PRAGMA table_info({table})")}
            if col.split()[0] not in existing:
                c.execute(f"ALTER TABLE {table} ADD COLUMN {col} {typ}".replace(f"{col} {typ}",f"{col.split()[0]} {typ}") if " " in col else f"ALTER TABLE {table} ADD COLUMN {col} {typ}")
                if backfill and backfill in existing:
                    c.execute(f"UPDATE {table} SET {col.split()[0]}={backfill}")
    except Exception:
        pass  # column may already exist with different default

# ══════════════════════════════════════════════════════
# TEXT UTILITIES
# ══════════════════════════════════════════════════════

_RE_FENCE  = re.compile(r"```[\s\S]*?```", re.M)
_RE_IMG    = re.compile(r"!\[[^\]]*\]\([^)]+\)")
_RE_LINK   = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_RE_HTML   = re.compile(r"<[^>]+>")
_RE_HEAD   = re.compile(r"^#{1,6}\s+", re.M)
_RE_BI     = re.compile(r"\*{1,3}(.+?)\*{1,3}", re.DOTALL)
_RE_STRIKE = re.compile(r"~~(.+?)~~", re.DOTALL)
_RE_CODE   = re.compile(r"`+(.+?)`+", re.DOTALL)
_RE_HR     = re.compile(r"^\s*[-*_]{3,}\s*$", re.M)
_RE_BQ     = re.compile(r"^>\s?", re.M)
_RE_LIST   = re.compile(r"^\s*[-*+]\s+|^\s*\d+\.\s+", re.M)

def strip_md(text: str) -> str:
    if not text: return ""
    text = _RE_FENCE.sub(" ", text)
    text = _RE_IMG.sub("", text)
    text = _RE_LINK.sub(r"\1", text)
    text = _RE_HTML.sub("", text)
    text = _RE_HEAD.sub("", text)
    text = _RE_BI.sub(r"\1", text)
    text = _RE_STRIKE.sub(r"\1", text)
    text = _RE_CODE.sub(r"\1", text)
    text = _RE_HR.sub("", text)
    text = _RE_BQ.sub("", text)
    text = _RE_LIST.sub("", text)
    return re.sub(r"  +", " ", text).strip()

def count_words(text: str) -> int:
    if not text or not text.strip(): return 0
    return len(strip_md(text).split())

def reading_time(word_count: int) -> str:
    """Estimate reading time at 238 words/min (average adult reader)."""
    mins = max(1, round(word_count / 238))
    if mins < 60: return f"{mins} min"
    h = mins // 60; m = mins % 60
    return f"{h}h{m:02d}min" if m else f"{h}h"

def writing_time(word_count: int) -> str:
    """Estimate writing time at 40 words/min (average typing speed for prose)."""
    mins = max(1, round(word_count / 40))
    if mins < 60: return f"{mins} min"
    h = mins // 60; m = mins % 60
    return f"{h}h{m:02d}min" if m else f"{h}h"

# ══════════════════════════════════════════════════════
# REPOSITORIES
# ══════════════════════════════════════════════════════

def _rows(c, sql, params=()):
    return [dict(r) for r in c.execute(sql, params).fetchall()]

def _one(c, sql, params=()):
    r = c.execute(sql, params).fetchone()
    return dict(r) if r else None

# ── Books ──────────────────────────────────────────────

def book_create(title, synopsis="", genre="", status="Planejamento", img=None, mime="image/jpeg"):
    with _conn() as c:
        return c.execute(
            "INSERT INTO books(title,synopsis,genre,status,cover_image,cover_mime) VALUES(?,?,?,?,?,?)",
            (title, synopsis, genre, status, img, mime)).lastrowid

def book_update(bid, **kw):
    allowed = {"title","synopsis","genre","status","cover_image","cover_mime","word_count","last_opened"}
    data = {k:v for k,v in kw.items() if k in allowed}
    if not data: return
    data["updated_at"] = "CURRENT_TIMESTAMP"
    sets = ",".join(f"{k}=CURRENT_TIMESTAMP" if v=="CURRENT_TIMESTAMP" else f"{k}=?" for k,v in data.items())
    vals = [v for v in data.values() if v!="CURRENT_TIMESTAMP"] + [bid]
    with _conn() as c: c.execute(f"UPDATE books SET {sets} WHERE id=?", vals)

def book_touch(bid):
    """Update last_opened timestamp when user opens a book."""
    with _conn() as c: c.execute("UPDATE books SET last_opened=CURRENT_TIMESTAMP WHERE id=?", (bid,))

def book_soft_delete(bid):
    with _conn() as c: c.execute("UPDATE books SET deleted_at=CURRENT_TIMESTAMP WHERE id=? AND deleted_at IS NULL", (bid,))

def book_restore(bid):
    with _conn() as c: c.execute("UPDATE books SET deleted_at=NULL WHERE id=?", (bid,))

def book_hard_delete(bid):
    with _conn() as c: c.execute("DELETE FROM books WHERE id=?", (bid,))

def book_get(bid):
    with _conn() as c: return _one(c, "SELECT * FROM books WHERE id=? AND deleted_at IS NULL", (bid,))

def book_list(query="", genre=None, status=None, sort="updated"):
    sql = "SELECT * FROM books WHERE deleted_at IS NULL"
    p = []
    if query: sql += " AND(title LIKE? OR synopsis LIKE?)"; p += [f"%{query}%"]*2
    if genre: sql += " AND genre=?"; p.append(genre)
    if status: sql += " AND status=?"; p.append(status)
    order = {"updated":"updated_at DESC","created":"created_at DESC","title":"title ASC","opened":"COALESCE(last_opened,created_at) DESC"}.get(sort,"updated_at DESC")
    sql += f" ORDER BY {order}"
    with _conn() as c: return _rows(c, sql, p)

def book_list_lw(sort="updated"):
    order = {"updated":"updated_at DESC","created":"created_at DESC","title":"title ASC","opened":"COALESCE(last_opened,created_at) DESC"}.get(sort,"updated_at DESC")
    with _conn() as c:
        return _rows(c, f"SELECT id,title,synopsis,genre,status,word_count,created_at,updated_at,last_opened FROM books WHERE deleted_at IS NULL ORDER BY {order}")

def book_list_deleted():
    with _conn() as c:
        return _rows(c, "SELECT id,title,genre,status,word_count,deleted_at FROM books WHERE deleted_at IS NOT NULL ORDER BY deleted_at DESC")

def book_genres():
    with _conn() as c:
        return [r[0] for r in c.execute("SELECT DISTINCT genre FROM books WHERE genre!='' AND deleted_at IS NULL ORDER BY genre")]

def book_recalc_wc(bid):
    with _conn() as c:
        total = c.execute("SELECT COALESCE(SUM(word_count),0) FROM chapters WHERE book_id=? AND deleted_at IS NULL", (bid,)).fetchone()[0]
        c.execute("UPDATE books SET word_count=?,updated_at=CURRENT_TIMESTAMP WHERE id=?", (total, bid))
    return total

def book_stats():
    with _conn() as c:
        tb = c.execute("SELECT COUNT(*) FROM books WHERE deleted_at IS NULL").fetchone()[0]
        tc = c.execute("SELECT COUNT(*) FROM chapters WHERE deleted_at IS NULL").fetchone()[0]
        tw = c.execute("SELECT COALESCE(SUM(word_count),0) FROM books WHERE deleted_at IS NULL").fetchone()[0]
        bs = _rows(c, "SELECT status,COUNT(*) as n FROM books WHERE deleted_at IS NULL GROUP BY status")
        streak = _calc_streak(c)
    return {"total_books":tb,"total_chapters":tc,"total_words":tw,
            "by_status":{r["status"]:r["n"] for r in bs},"streak":streak}

def _calc_streak(c):
    """Calculate current writing streak in days."""
    rows = c.execute(
        "SELECT DISTINCT session_date FROM writing_sessions "
        "WHERE words_written>0 ORDER BY session_date DESC LIMIT 365"
    ).fetchall()
    if not rows: return 0
    from datetime import date, timedelta
    dates = {r[0] for r in rows}
    today = date.today().isoformat()
    yesterday = (date.today() - timedelta(days=1)).isoformat()
    if today not in dates and yesterday not in dates: return 0
    streak = 0
    check = date.today() if today in dates else date.today() - timedelta(days=1)
    while check.isoformat() in dates:
        streak += 1
        check -= timedelta(days=1)
    return streak

# ── Chapters ───────────────────────────────────────────

def ch_create(book_id, title, content=""):
    wc = count_words(content)
    with _conn() as c:
        c.execute("BEGIN IMMEDIATE")
        pos = c.execute("SELECT COALESCE(MAX(position),-1)+1 FROM chapters WHERE book_id=? AND deleted_at IS NULL", (book_id,)).fetchone()[0]
        cid = c.execute("INSERT INTO chapters(book_id,title,content,position,word_count) VALUES(?,?,?,?,?)", (book_id,title,content,pos,wc)).lastrowid
    book_recalc_wc(book_id); return cid

def ch_get(cid):
    with _conn() as c: return _one(c, "SELECT * FROM chapters WHERE id=? AND deleted_at IS NULL", (cid,))

def ch_list(book_id):
    with _conn() as c: return _rows(c, "SELECT * FROM chapters WHERE book_id=? AND deleted_at IS NULL ORDER BY position", (book_id,))

def ch_list_lw(book_id):
    with _conn() as c:
        return _rows(c, "SELECT id,book_id,title,position,word_count,word_goal,scene_marker,created_at,updated_at FROM chapters WHERE book_id=? AND deleted_at IS NULL ORDER BY position", (book_id,))

def ch_list_deleted(book_id):
    with _conn() as c: return _rows(c, "SELECT id,title,word_count,deleted_at FROM chapters WHERE book_id=? AND deleted_at IS NOT NULL ORDER BY deleted_at DESC", (book_id,))

def ch_get_wc(cid):
    with _conn() as c:
        r = c.execute("SELECT word_count FROM chapters WHERE id=? AND deleted_at IS NULL", (cid,)).fetchone()
        return r["word_count"] if r else 0

def ch_save(cid, content, book_id, minutes=0):
    prev = ch_get_wc(cid); wc = count_words(content)
    with _conn() as c:
        c.execute("UPDATE chapters SET content=?,word_count=?,updated_at=CURRENT_TIMESTAMP WHERE id=? AND deleted_at IS NULL", (content,wc,cid))
    book_recalc_wc(book_id)
    delta = wc - prev
    if delta > 0: sess_log(delta, book_id, cid, minutes)

def ch_update(cid, **kw):
    allowed = {"title","position","word_goal","scene_marker"}
    data = {k:v for k,v in kw.items() if k in allowed}
    if not data: return
    data["updated_at"] = "CURRENT_TIMESTAMP"
    sets = ",".join(f"{k}=CURRENT_TIMESTAMP" if v=="CURRENT_TIMESTAMP" else f"{k}=?" for k,v in data.items())
    vals = [v for v in data.values() if v!="CURRENT_TIMESTAMP"] + [cid]
    with _conn() as c: c.execute(f"UPDATE chapters SET {sets} WHERE id=?", vals)

def ch_reorder(book_id, ids):
    with _conn() as c:
        for pos,cid in enumerate(ids):
            c.execute("UPDATE chapters SET position=? WHERE id=? AND book_id=?", (pos,cid,book_id))

def ch_soft_delete(cid, book_id):
    with _conn() as c: c.execute("UPDATE chapters SET deleted_at=CURRENT_TIMESTAMP WHERE id=? AND deleted_at IS NULL", (cid,))
    book_recalc_wc(book_id)

def ch_restore(cid, book_id):
    with _conn() as c: c.execute("UPDATE chapters SET deleted_at=NULL WHERE id=?", (cid,))
    book_recalc_wc(book_id)

def ch_hard_delete(cid, book_id):
    with _conn() as c: c.execute("DELETE FROM chapters WHERE id=?", (cid,))
    book_recalc_wc(book_id)

def ch_stats_for_book(book_id):
    """Return per-chapter stats for the book overview."""
    with _conn() as c:
        rows = _rows(c, "SELECT id,title,word_count,word_goal,scene_marker,updated_at FROM chapters WHERE book_id=? AND deleted_at IS NULL ORDER BY position", (book_id,))
    for r in rows:
        r["reading_time"] = reading_time(r["word_count"])
        r["goal_pct"] = min(100, int(r["word_count"]/r["word_goal"]*100)) if r.get("word_goal") else 0
    return rows

# ── Characters ─────────────────────────────────────────

def char_create(book_id, name, role="", desc="", photo=None, mime="image/jpeg", rel="", notes="", age=""):
    with _conn() as c:
        return c.execute(
            "INSERT INTO characters(book_id,name,role,description,photo,photo_mime,relationships,notes,age) VALUES(?,?,?,?,?,?,?,?,?)",
            (book_id,name,role,desc,photo,mime,rel,notes,age)).lastrowid

def char_list(book_id, role_filter=None):
    sql = "SELECT * FROM characters WHERE book_id=? AND deleted_at IS NULL"
    p = [book_id]
    if role_filter: sql += " AND role=?"; p.append(role_filter)
    sql += " ORDER BY name"
    with _conn() as c: return _rows(c, sql, p)

def char_get(cid):
    with _conn() as c: return _one(c, "SELECT * FROM characters WHERE id=? AND deleted_at IS NULL", (cid,))

def char_update(cid, **kw):
    allowed = {"name","role","description","photo","photo_mime","relationships","notes","age"}
    data = {k:v for k,v in kw.items() if k in allowed}
    if not data: return
    data["updated_at"] = "CURRENT_TIMESTAMP"
    sets = ",".join(f"{k}=CURRENT_TIMESTAMP" if v=="CURRENT_TIMESTAMP" else f"{k}=?" for k,v in data.items())
    vals = [v for v in data.values() if v!="CURRENT_TIMESTAMP"] + [cid]
    with _conn() as c: c.execute(f"UPDATE characters SET {sets} WHERE id=?", vals)

def char_soft_delete(cid):
    with _conn() as c: c.execute("UPDATE characters SET deleted_at=CURRENT_TIMESTAMP WHERE id=?", (cid,))

def char_search(book_id, q):
    with _conn() as c:
        return _rows(c, "SELECT * FROM characters WHERE book_id=? AND deleted_at IS NULL AND(name LIKE? OR description LIKE? OR role LIKE? OR notes LIKE?) ORDER BY name",
                     (book_id,f"%{q}%",f"%{q}%",f"%{q}%",f"%{q}%"))

def char_roles(book_id):
    with _conn() as c:
        return [r[0] for r in c.execute("SELECT DISTINCT role FROM characters WHERE book_id=? AND role!='' AND deleted_at IS NULL ORDER BY role", (book_id,))]

# ── World Building ─────────────────────────────────────

def loc_create(book_id, name, desc="", notes=""):
    with _conn() as c: return c.execute("INSERT INTO locations(book_id,name,description,notes) VALUES(?,?,?,?)", (book_id,name,desc,notes)).lastrowid

def loc_list(book_id):
    with _conn() as c: return _rows(c, "SELECT * FROM locations WHERE book_id=? AND deleted_at IS NULL ORDER BY name", (book_id,))

def loc_update(lid, **kw):
    allowed = {"name","description","notes"}
    data = {k:v for k,v in kw.items() if k in allowed}
    if not data: return
    data["updated_at"] = "CURRENT_TIMESTAMP"
    sets = ",".join(f"{k}=CURRENT_TIMESTAMP" if v=="CURRENT_TIMESTAMP" else f"{k}=?" for k,v in data.items())
    vals = [v for v in data.values() if v!="CURRENT_TIMESTAMP"] + [lid]
    with _conn() as c: c.execute(f"UPDATE locations SET {sets} WHERE id=?", vals)

def loc_soft_delete(lid):
    with _conn() as c: c.execute("UPDATE locations SET deleted_at=CURRENT_TIMESTAMP WHERE id=?", (lid,))

def loc_search(book_id, q):
    with _conn() as c:
        return _rows(c, "SELECT * FROM locations WHERE book_id=? AND deleted_at IS NULL AND(name LIKE? OR description LIKE?) ORDER BY name", (book_id,f"%{q}%",f"%{q}%"))

def fac_create(book_id, name, desc="", notes=""):
    with _conn() as c: return c.execute("INSERT INTO factions(book_id,name,description,notes) VALUES(?,?,?,?)", (book_id,name,desc,notes)).lastrowid

def fac_list(book_id):
    with _conn() as c: return _rows(c, "SELECT * FROM factions WHERE book_id=? AND deleted_at IS NULL ORDER BY name", (book_id,))

def fac_update(fid, **kw):
    allowed = {"name","description","notes"}
    data = {k:v for k,v in kw.items() if k in allowed}
    if not data: return
    data["updated_at"] = "CURRENT_TIMESTAMP"
    sets = ",".join(f"{k}=CURRENT_TIMESTAMP" if v=="CURRENT_TIMESTAMP" else f"{k}=?" for k,v in data.items())
    vals = [v for v in data.values() if v!="CURRENT_TIMESTAMP"] + [fid]
    with _conn() as c: c.execute(f"UPDATE factions SET {sets} WHERE id=?", vals)

def fac_soft_delete(fid):
    with _conn() as c: c.execute("UPDATE factions SET deleted_at=CURRENT_TIMESTAMP WHERE id=?", (fid,))

def ev_create(book_id, title, desc="", date_label=""):
    with _conn() as c:
        c.execute("BEGIN IMMEDIATE")
        pos = c.execute("SELECT COALESCE(MAX(position),-1)+1 FROM timeline_events WHERE book_id=? AND deleted_at IS NULL", (book_id,)).fetchone()[0]
        return c.execute("INSERT INTO timeline_events(book_id,title,description,date_label,position) VALUES(?,?,?,?,?)", (book_id,title,desc,date_label,pos)).lastrowid

def ev_list(book_id):
    with _conn() as c: return _rows(c, "SELECT * FROM timeline_events WHERE book_id=? AND deleted_at IS NULL ORDER BY position", (book_id,))

def ev_update(eid, **kw):
    allowed = {"title","description","date_label","position"}
    data = {k:v for k,v in kw.items() if k in allowed}
    if not data: return
    data["updated_at"] = "CURRENT_TIMESTAMP"
    sets = ",".join(f"{k}=CURRENT_TIMESTAMP" if v=="CURRENT_TIMESTAMP" else f"{k}=?" for k,v in data.items())
    vals = [v for v in data.values() if v!="CURRENT_TIMESTAMP"] + [eid]
    with _conn() as c: c.execute(f"UPDATE timeline_events SET {sets} WHERE id=?", vals)

def ev_soft_delete(eid):
    with _conn() as c: c.execute("UPDATE timeline_events SET deleted_at=CURRENT_TIMESTAMP WHERE id=?", (eid,))

# ── Brain Dump ─────────────────────────────────────────

def bd_create(content, book_id=None, tags=""):
    with _conn() as c: return c.execute("INSERT INTO brain_dumps(content,book_id,tags) VALUES(?,?,?)", (content,book_id,tags)).lastrowid

def bd_list(query="", book_id=None, tag=None, pinned_first=True):
    sql = "SELECT * FROM brain_dumps WHERE deleted_at IS NULL"
    p = []
    if book_id: sql += " AND book_id=?"; p.append(book_id)
    if query: sql += " AND(content LIKE? OR tags LIKE?)"; p += [f"%{query}%"]*2
    if tag: sql += " AND tags LIKE?"; p.append(f"%{tag}%")
    sql += " ORDER BY pinned DESC, created_at DESC" if pinned_first else " ORDER BY created_at DESC"
    with _conn() as c: return _rows(c, sql, p)

def bd_update(did, content, tags):
    with _conn() as c: c.execute("UPDATE brain_dumps SET content=?,tags=?,updated_at=CURRENT_TIMESTAMP WHERE id=?", (content,tags,did))

def bd_toggle_pin(did):
    with _conn() as c:
        cur = c.execute("SELECT pinned FROM brain_dumps WHERE id=?", (did,)).fetchone()
        if cur: c.execute("UPDATE brain_dumps SET pinned=? WHERE id=?", (0 if cur["pinned"] else 1, did))

def bd_soft_delete(did):
    with _conn() as c: c.execute("UPDATE brain_dumps SET deleted_at=CURRENT_TIMESTAMP WHERE id=?", (did,))

def bd_tags():
    with _conn() as c: rows = _rows(c, "SELECT tags FROM brain_dumps WHERE tags!='' AND deleted_at IS NULL")
    tags = set()
    for r in rows:
        for t in r["tags"].split(","):
            t = t.strip()
            if t: tags.add(t)
    return sorted(tags)

# ── Goals & Sessions ───────────────────────────────────

def goal_save(target, period):
    with _conn() as c:
        c.execute("UPDATE goals SET active=0 WHERE period=?", (period,))
        c.execute("INSERT INTO goals(goal_type,target_words,period,active) VALUES('writing',?,?,1)", (target,period))

def goal_active():
    with _conn() as c: return _rows(c, "SELECT * FROM goals WHERE active=1")

def sess_log(words, book_id=None, chapter_id=None, minutes=0):
    with _conn() as c: c.execute("INSERT INTO writing_sessions(book_id,chapter_id,words_written,minutes_written) VALUES(?,?,?,?)", (book_id,chapter_id,words,minutes))

def sess_today():
    with _conn() as c: return c.execute("SELECT COALESCE(SUM(words_written),0) FROM writing_sessions WHERE session_date=DATE('now')").fetchone()[0]

def sess_month():
    with _conn() as c: return c.execute("SELECT COALESCE(SUM(words_written),0) FROM writing_sessions WHERE strftime('%Y-%m',session_date)=strftime('%Y-%m','now')").fetchone()[0]

def sess_30days():
    with _conn() as c: return _rows(c, "SELECT session_date,SUM(words_written) as total FROM writing_sessions WHERE session_date>=DATE('now','-30 days') GROUP BY session_date ORDER BY session_date")

def sess_heatmap():
    """Return 365-day heatmap data for the calendar view."""
    with _conn() as c:
        return _rows(c, "SELECT session_date,SUM(words_written) as total FROM writing_sessions WHERE session_date>=DATE('now','-364 days') GROUP BY session_date")

def sess_best_day():
    with _conn() as c:
        r = c.execute("SELECT session_date,SUM(words_written) as total FROM writing_sessions GROUP BY session_date ORDER BY total DESC LIMIT 1").fetchone()
        return dict(r) if r else {"session_date":"—","total":0}

def sess_total_minutes():
    with _conn() as c:
        r = c.execute("SELECT COALESCE(SUM(minutes_written),0) FROM writing_sessions").fetchone()
        return r[0] if r else 0

# ── Settings ───────────────────────────────────────────

def setting_get(key, default=None):
    with _conn() as c:
        r = c.execute("SELECT value FROM app_settings WHERE key=?", (key,)).fetchone()
        return r["value"] if r else default

def setting_set(key, value):
    with _conn() as c: c.execute("INSERT OR REPLACE INTO app_settings(key,value) VALUES(?,?)", (key,value))

def kindle_pos_get():
    bid = setting_get("kindle_book_id"); idx = setting_get("kindle_chapter_idx","0")
    return (int(bid) if bid else None), int(idx)

def kindle_pos_save(book_id, idx):
    setting_set("kindle_book_id", str(book_id)); setting_set("kindle_chapter_idx", str(idx))

# ── Kindle bookmarks ───────────────────────────────────

def bookmark_add(book_id, chapter_id, note=""):
    with _conn() as c: return c.execute("INSERT INTO kindle_bookmarks(book_id,chapter_id,note) VALUES(?,?,?)", (book_id,chapter_id,note)).lastrowid

def bookmark_list(book_id):
    with _conn() as c: return _rows(c, "SELECT kb.*,ch.title as ch_title FROM kindle_bookmarks kb LEFT JOIN chapters ch ON ch.id=kb.chapter_id WHERE kb.book_id=? ORDER BY kb.created_at DESC", (book_id,))

def bookmark_delete(bkid):
    with _conn() as c: c.execute("DELETE FROM kindle_bookmarks WHERE id=?", (bkid,))

# ══════════════════════════════════════════════════════
# IMAGE UTILITIES
# ══════════════════════════════════════════════════════

def img_process(f, max_size=(400,600)):
    img = Image.open(f); img.thumbnail(max_size, Image.LANCZOS)
    if img.mode in ("RGBA","P"): img = img.convert("RGB")
    buf = io.BytesIO(); img.save(buf, format="JPEG", quality=85)
    return buf.getvalue(), "image/jpeg"

def img_b64(data):
    if not data: return ""
    return f"data:image/jpeg;base64,{base64.b64encode(data).decode()}"

# ══════════════════════════════════════════════════════
# EXPORT
# ══════════════════════════════════════════════════════

def export_pdf(book_id):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable, Image
    book = book_get(book_id); chapters = ch_list(book_id)
    buf = io.BytesIO(); W,H = A4
    def footer(canvas, doc):
        canvas.saveState(); canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.HexColor("#999999"))
        canvas.drawCentredString(W/2, 1.5*cm, str(canvas.getPageNumber())); canvas.restoreState()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=3*cm, leftMargin=3*cm, topMargin=3*cm, bottomMargin=2.5*cm, onFirstPage=footer, onLaterPages=footer)
    S = getSampleStyleSheet()
    title_s = ParagraphStyle("T", parent=S["Title"], fontSize=28, textColor=colors.HexColor("#1a1a2e"), alignment=TA_CENTER)
    sub_s   = ParagraphStyle("S", parent=S["Normal"], fontSize=13, textColor=colors.HexColor("#555"), alignment=TA_CENTER)
    ch_s    = ParagraphStyle("C", parent=S["Heading1"], fontSize=20, spaceAfter=12)
    body_s  = ParagraphStyle("B", parent=S["Normal"], fontSize=12, leading=19, spaceAfter=8, alignment=TA_JUSTIFY, firstLineIndent=18)
    toc_s   = ParagraphStyle("TC", parent=S["Normal"], fontSize=11, leading=16, spaceAfter=4, leftIndent=12)
    story = []
    if book.get("cover_image"):
        try:
            img = Image(io.BytesIO(book["cover_image"]), width=10*cm, height=15*cm, kind="proportional")
            img.hAlign = "CENTER"; story += [Spacer(1,cm), img, Spacer(1,cm)]
        except: pass
    story.append(Paragraph(book["title"], title_s))
    if book.get("genre"): story.append(Paragraph(f"Gênero: {book['genre']}", sub_s))
    if book.get("synopsis"): story += [Spacer(1,.5*cm), Paragraph(book["synopsis"], body_s)]
    story.append(PageBreak())
    story.append(Paragraph("Sumário", ParagraphStyle("TT", parent=S["Heading1"], fontSize=18, spaceAfter=12)))
    story.append(HRFlowable(width="100%", thickness=.5, color=colors.HexColor("#ccc")))
    story.append(Spacer(1,.3*cm))
    for i,ch in enumerate(chapters,1): story.append(Paragraph(f"{i}. {ch['title']}", toc_s))
    story.append(PageBreak())
    for ch in chapters:
        story.append(Paragraph(ch["title"], ch_s))
        story.append(HRFlowable(width="100%", thickness=.5, color=colors.HexColor("#ddd")))
        story.append(Spacer(1,.3*cm))
        for para in strip_md(ch.get("content") or "").split("\n\n"):
            para = para.strip()
            if para: story.append(Paragraph(para.replace("\n","<br/>"), body_s))
        story.append(PageBreak())
    doc.build(story); return buf.getvalue()

def export_docx(book_id):
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    book = book_get(book_id); chapters = ch_list(book_id)
    doc = Document()
    if book.get("cover_image"):
        try:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(book["cover_image"]), width=Inches(3.5))
            doc.add_paragraph()
        except: pass
    h = doc.add_heading(book["title"], 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if book.get("genre"):
        p = doc.add_paragraph(f"Gênero: {book['genre']}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if book.get("synopsis"):
        doc.add_paragraph(); doc.add_heading("Sinopse", level=2); doc.add_paragraph(book["synopsis"])
    doc.add_page_break()
    doc.add_heading("Sumário", level=1)
    for i,ch in enumerate(chapters,1): doc.add_paragraph(f"{i}. {ch['title']}")
    doc.add_page_break()
    for ch in chapters:
        doc.add_heading(ch["title"], level=1)
        for para in strip_md(ch.get("content") or "").split("\n\n"):
            para = para.strip()
            if para: doc.add_paragraph(para)
        doc.add_page_break()
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

def export_epub(book_id):
    from ebooklib import epub
    import markdown2
    book = book_get(book_id); chapters = ch_list(book_id)
    eb = epub.EpubBook(); eb.set_title(book["title"]); eb.set_language("pt")
    eb.add_metadata("DC","description", book.get("synopsis") or "")
    if book.get("cover_image"):
        try: eb.set_cover("cover.jpg", book["cover_image"])
        except: pass
    css = epub.EpubItem(uid="style", file_name="style.css", media_type="text/css",
        content="body{font-family:Georgia,serif;font-size:1em;line-height:1.7;margin:1.5em 2em}h1{font-size:1.9em;text-align:center;border-bottom:1px solid #ddd;padding-bottom:.5em;margin-bottom:1em}p{margin:0 0 .8em;text-align:justify;text-indent:1.5em}p:first-of-type{text-indent:0}")
    eb.add_item(css)
    ep_chs = []; toc = []
    for i,ch in enumerate(chapters,1):
        ec = epub.EpubHtml(title=ch["title"], file_name=f"ch{i:03d}.xhtml", lang="pt")
        ec.content = f"<html><body><h1>{ch['title']}</h1>{markdown2.markdown(ch.get('content') or '')}</body></html>"
        ec.add_item(css); eb.add_item(ec); ep_chs.append(ec)
        toc.append(epub.Link(f"ch{i:03d}.xhtml", ch["title"], f"ch{i}"))
    eb.toc = toc; eb.add_item(epub.EpubNcx()); eb.add_item(epub.EpubNav())
    eb.spine = ["nav"] + ep_chs
    buf = io.BytesIO(); epub.write_epub(buf, eb); return buf.getvalue()

# ── Motivational quotes (offline) ──────────────────────

QUOTES = [
    ("A primeira versão é apenas você contando a história para si mesma.", "Terry Pratchett"),
    ("Escreva bêbada, revise sóbria.", "Ernest Hemingway"),
    ("Você não tem que ser boa para começar, mas tem que começar para ser boa.", "Joe Sabah"),
    ("Uma palavra depois da outra. É assim que se escreve um livro.", "Neil Gaiman"),
    ("Não existe uma musa. Existe apenas escrever.", "Hemingway"),
    ("O talento é mais barato que o sal. O que separa o artista talentoso do bem-sucedido é trabalho, trabalho e mais trabalho.", "Stephen King"),
    ("Escrever é a única forma que tenho de orar.", "Cormac McCarthy"),
    ("Comece onde você está. Use o que você tem. Faça o que você pode.", "Arthur Ashe"),
    ("O primeiro rascunho de qualquer coisa é uma merda.", "Hemingway"),
    ("Leia muito, escreva muito.", "Stephen King"),
    ("Palavras são nossas maiores fonte de magia e prejuízo ao mesmo tempo.", "J.K. Rowling"),
    ("A coragem de escrever é a coragem de se revelar.", "May Sarton"),
]

def daily_quote():
    from datetime import date
    idx = date.today().toordinal() % len(QUOTES)
    return QUOTES[idx]
